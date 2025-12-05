"""
Document processing utilities for .docx files.
Handles reading and exporting documents with accepted edits.
"""

import io
from dataclasses import dataclass
from typing import BinaryIO

from docx import Document


@dataclass
class AcceptedEdit:
    """Represents an accepted edit with optional custom text."""

    para_index: int
    edit_index: int
    original_text: str
    final_text: str  # Either the AI's suggestion or custom user text


@dataclass
class ProcessedParagraph:
    """Holds a paragraph's text and its index in the document."""

    index: int
    text: str
    original_paragraph: object  # docx Paragraph object


class DocumentProcessor:
    """
    Handles reading and writing .docx files.
    """

    def __init__(self):
        self.document: Document | None = None
        self.paragraphs: list[ProcessedParagraph] = []

    def load_from_bytes(self, file_bytes: BinaryIO) -> list[str]:
        """
        Load a document from file bytes (e.g., from Streamlit uploader).

        Args:
            file_bytes: File-like object containing .docx data

        Returns:
            List of paragraph texts
        """
        self.document = Document(file_bytes)
        self.paragraphs = []

        for i, para in enumerate(self.document.paragraphs):
            self.paragraphs.append(
                ProcessedParagraph(index=i, text=para.text, original_paragraph=para)
            )

        return [p.text for p in self.paragraphs]

    def load_from_path(self, path: str) -> list[str]:
        """
        Load a document from a file path.

        Args:
            path: Path to the .docx file

        Returns:
            List of paragraph texts
        """
        self.document = Document(path)
        self.paragraphs = []

        for i, para in enumerate(self.document.paragraphs):
            self.paragraphs.append(
                ProcessedParagraph(index=i, text=para.text, original_paragraph=para)
            )

        return [p.text for p in self.paragraphs]

    def export_with_accepted_edits(self, accepted_edits: list[AcceptedEdit]) -> bytes:
        """
        Export document with only the accepted edits applied.

        Args:
            accepted_edits: List of AcceptedEdit objects containing the edits to apply

        Returns:
            Bytes of the modified .docx file
        """
        if not self.document:
            raise ValueError("No document loaded. Call load_from_bytes or load_from_path first.")

        # Group edits by paragraph index
        edits_by_para: dict[int, list[AcceptedEdit]] = {}
        for edit in accepted_edits:
            if edit.para_index not in edits_by_para:
                edits_by_para[edit.para_index] = []
            edits_by_para[edit.para_index].append(edit)

        # Create a new document
        new_doc = Document()

        for i, proc_para in enumerate(self.paragraphs):
            # Start with original text
            new_text = proc_para.text

            # Apply accepted edits for this paragraph
            if i in edits_by_para:
                para_edits = edits_by_para[i]
                # Sort by length (longest first) to avoid partial replacement issues
                para_edits_sorted = sorted(
                    para_edits, key=lambda e: len(e.original_text), reverse=True
                )
                for edit in para_edits_sorted:
                    new_text = new_text.replace(edit.original_text, edit.final_text, 1)

            # Add paragraph to new document
            new_para = new_doc.add_paragraph()
            self._copy_paragraph_format(proc_para.original_paragraph, new_para)
            new_para.add_run(new_text)

        # Save to bytes
        buffer = io.BytesIO()
        new_doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _copy_paragraph_format(self, source_para, target_para):
        """Copy paragraph formatting from source to target."""
        try:
            if source_para.style:
                target_para.style = source_para.style.name
        except Exception:
            pass  # Style may not exist in new document
